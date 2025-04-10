import math

from docx import Document
from docx.enum.text import WD_BREAK
import random
from docx.shared import Pt
import datetime

first_names = ["Rachanon", "Pat", "Suthipong", "Thapanat", "Nattawut", "Nat", "Wichayaporn", "Natniti", "Krititach",
               "Tiya", "Thawan", "Opaspong", "Nichanan", "Daranporn", "Surasak", "Narit", "Jaipak", "Sajaporn",
               "Piyaporn", "Tanisra", "Sasima", "Sunisorn", "Pawin", "Setha", "Petcharat", "Chama", "Kongkiat",
               "Bulakorn", "Narot", "Sorat", "Pongdech", "Dolawa", "Supitcha", "Chidchanok", "Tanisorn", "Nawisa",
               "Rat", "Trin", "Possathorn", "Nawaporn", "Jakkapat", "Weerachai", "Pakpoom", "Pachariya", "Sira",
               "Janwipa", "Nipa", "Prisana", "Phurinat", "Pariya", "Ptchaya", "Supreeya", "Rachanon", "Pat",
               "Suthipong", "Thapanat", "Nattawut", "Nat", "Wichayaporn", "Natniti", "Krititach",
               "Tiya", "Thawan", "Opaspong", "Nichanan", "Daranporn", "Surasak", "Narit", "Jaipak", "Sajaporn",
               "Piyaporn", "Tanisra", "Sasima", "Sunisorn", "Pawin", "Setha", "Petcharat", "Chama", "Kongkiat",
               "Bulakorn", "Narot", "Sorat", "Pongdech", "Dolawa", "Supitcha", "Chidchanok", "Tanisorn", "Nawisa",
               "Rat", "Trin", "Possathorn", "Nawaporn", "Jakkapat", "Weerachai", "Pakpoom", "Pachariya", "Sira",
               "Janwipa", "Nipa", "Prisana", "Phurinat", "Pariya", "Ptchaya", "Supreeya", "Rachanon", "Pat",
               "Suthipong", "Thapanat", "Nattawut", "Nat", "Wichayaporn", "Natniti", "Krititach",
               "Tiya", "Thawan", "Opaspong", "Nichanan", "Daranporn", "Surasak", "Narit", "Jaipak", "Sajaporn",
               "Piyaporn", "Tanisra", "Sasima", "Sunisorn", "Pawin", "Setha", "Petcharat", "Chama", "Kongkiat",
               "Bulakorn", "Narot", "Sorat", "Pongdech", "Dolawa", "Supitcha", "Chidchanok", "Tanisorn", "Nawisa",
               "Rat", "Trin", "Possathorn", "Nawaporn", "Jakkapat", "Weerachai", "Pakpoom", "Pachariya", "Sira",
               "Janwipa", "Nipa", "Prisana", "Phurinat", "Pariya", "Ptchaya", "Supreeya"]
last_names = ["Tarawong", "Tongchai", "Satitayuth", "Santipipak", "Sawangwittaya", "Suwansiri", "Wongsawad",
              "Pleumpiti", "Suwannachok", "Udompak", "Jaroenpolwattana", "Lertwittaya", "Ngamkwan", "Satitayuth",
              "Roongrudee", "Siripachasab", "Pongpinij", "Worapaksirisakul", "Pititat", "Kwankaew", "Sukseree",
              "Jaroensakul", "Reongrit", "Lertkoonawong", "Achachai", "Sabtara", "Sangsuwan", "Pongwiroj", "Sritong",
              "Intaraprasart", "Thanakorn", "Rattanasetha", "Jaroenpolwattana", "Koomwong", "Arunrung", "Chokchai",
              "Pongpaisan", "Pongpattana", "Lertwittaya", "Treewuth", "Srisoi", "Paisarn", "Jantarasab", "Makmee",
              "Jarastham", "Tantakarun", "Sukseree", "Yinggnam", "Siritada", "Pleumpiti", "Udomsab", "Sirikoson",
              "Tarawong", "Tongchai", "Satitayuth", "Santipipak", "Sawangwittaya", "Suwansiri", "Wongsawad",
              "Pleumpiti", "Suwannachok", "Udompak", "Jaroenpolwattana", "Lertwittaya", "Ngamkwan", "Satitayuth",
              "Roongrudee", "Siripachasab", "Pongpinij", "Worapaksirisakul", "Pititat", "Kwankaew", "Sukseree",
              "Jaroensakul", "Reongrit", "Lertkoonawong", "Achachai", "Sabtara", "Sangsuwan", "Pongwiroj", "Sritong",
              "Intaraprasart", "Thanakorn", "Rattanasetha", "Jaroenpolwattana", "Koomwong", "Arunrung", "Chokchai",
              "Pongpaisan", "Pongpattana", "Lertwittaya", "Treewuth", "Srisoi", "Paisarn", "Jantarasab", "Makmee",
              "Jarastham", "Tantakarun", "Sukseree", "Yinggnam", "Siritada", "Pleumpiti", "Udomsab", "Sirikoson",
              "Tarawong", "Tongchai", "Satitayuth", "Santipipak", "Sawangwittaya", "Suwansiri", "Wongsawad",
              "Pleumpiti", "Suwannachok", "Udompak", "Jaroenpolwattana", "Lertwittaya", "Ngamkwan", "Satitayuth",
              "Roongrudee", "Siripachasab", "Pongpinij", "Worapaksirisakul", "Pititat", "Kwankaew", "Sukseree",
              "Jaroensakul", "Reongrit", "Lertkoonawong", "Achachai", "Sabtara", "Sangsuwan", "Pongwiroj", "Sritong",
              "Intaraprasart", "Thanakorn", "Rattanasetha", "Jaroenpolwattana", "Koomwong", "Arunrung", "Chokchai",
              "Pongpaisan", "Pongpattana", "Lertwittaya", "Treewuth", "Srisoi", "Paisarn", "Jantarasab", "Makmee",
              "Jarastham", "Tantakarun", "Sukseree", "Yinggnam", "Siritada", "Pleumpiti", "Udomsab", "Sirikoson"]
positions = [("CEO", 1, 500_000),
             ("COO", 1, 300_000),
             ("CFO", 1, 300_000),
             ("CTO", 1, 300_000),
             ("CMO", 1, 300_000),
             ("director", 10, 100_000),
             ("team leader", 20, 50_000),
             ("senior", 30, 30_000),
             ("junior", 35, 15_000)]

company_names = ["Money Me Ltd.", "Account_ac Lyd.", "Lucky U Ltd.", "KWAM OD TON MEE CO., LTD.",
                 "P.B.R.N Corporation", "Umbrellar Corporation", "Bang Pha Pleun Company"]

random.seed(2023)

for name in company_names:
    fn = first_names.copy()
    ln = last_names.copy()
    random.shuffle(fn)
    random.shuffle(ln)
    document = Document()
    body_style = document.styles['Normal']
    font = body_style.font
    font.name = 'Times News Roman'
    font.size = Pt(16)
    Title_style = document.styles['Title']
    font = Title_style.font
    font.name = 'Times News Roman'
    font.size = Pt(18)
    font.bold = True
    all_data = [["Company", "id", "FirstName", "LastName", "Position", "Salary", "Height", "Weight", "DoY", "Year", "DoB"]]
    employee_id = 1
    for p in positions:
        for i in range(p[1] + random.randint(0, math.floor(p[1] / 2))):
            document.add_paragraph(f"{name}").style = Title_style
            data_row = [name, employee_id, fn.pop(), ln.pop(), p[0], math.floor(p[2] * (1 + random.random() / 3)),
                        random.randint(150, 190), random.randint(50, 90)]
            employee_id += 1
            day_of_year = random.randint(1, 365)
            birth_year = random.randint(1950, 2000)
            data_row.append(day_of_year)
            data_row.append(birth_year)
            data_row.append((datetime.datetime(birth_year, 1, 1) + datetime.timedelta(days=day_of_year - 1)).date())
            all_data.append(data_row)
            content = document.add_paragraph("First Name:\t")
            content.add_run(f"\t{data_row[2]}\t").underline = True
            content.add_run("\nLast Name: ")
            content.add_run(f"\t{data_row[3]}\t").underline = True
            content.add_run("\nPosition: ")
            content.add_run(f"\t{data_row[4]}\t").underline = True
            content.add_run("\nSalary: ")
            content.add_run(f"\t{data_row[5]}\t").underline = True
            content.add_run("\nHeight: ")
            content.add_run(f"\t{data_row[6]}\t").underline = True
            content.add_run("\nWeight: ")
            content.add_run(f"\t{data_row[7]}\t").underline = True
            content.add_run("\nDate of Birth: ")
            content.add_run(f"\t{data_row[9]}\t").underline = True
            content.add_run().add_break(WD_BREAK.PAGE)
    with open(f"employee_csv/{name}.txt", "w+") as txt_file:
        for line in all_data:
            line = [str(_) for _ in line]
            txt_file.write(",".join(line) + "\n")
    document.save(f"employee_docx/{name}.docx")
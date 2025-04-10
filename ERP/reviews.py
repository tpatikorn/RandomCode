import math

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
import random
from docx.shared import Pt, Inches
import datetime
import pandas as pd
from os import path, listdir

random.seed(2023)


def randint_plusminus(val_mid, val_min=1, val_max=5, val_range=2):
    return random.randint(max(val_mid % (val_max - val_min + 1) - val_range + val_min, val_min),
                          min(val_mid % (val_max - val_min + 1) + val_range + val_min, val_max))


for filename in listdir("generated_data/employee_csv"):
    company_name = filename[0: -4]
    data = pd.read_csv(path.join("generated_data/employee_csv", filename))
    junior_id = list(data[data.Position == "junior"].id)
    review_id = 1

    document = Document()
    body_style = document.styles['Normal']
    font = body_style.font
    font.name = 'Times News Roman'
    font.size = Pt(12)
    Title_style = document.styles['Title']
    font = Title_style.font
    font.name = 'Times News Roman'
    font.size = Pt(18)
    font.bold = True

    all_data = [["company_name", "review_id", "employee_id", "product_id",
                 "product_score_1", "product_score_2", "product_score_3",
                 "service_score_1", "service_score_2", "service_score_3",
                 "store_score_1", "store_score_2", "store_score_3", "overall_score"]]
    print(company_name)
    for i in range(1, random.randint(1000, 1100)):
        if i % 100 == 0:
            print(i)
        employee_id = random.sample(junior_id, 1)[0]
        product_id = random.randint(1, 20)
        product_score_1 = randint_plusminus(product_id)
        product_score_2 = randint_plusminus(product_id)
        product_score_3 = randint_plusminus(product_id)
        service_score_1 = randint_plusminus(employee_id)
        service_score_2 = randint_plusminus(employee_id)
        service_score_3 = randint_plusminus(employee_id)
        store_score_1 = random.randint(1, 5)
        store_score_2 = random.randint(1, 5)
        store_score_3 = random.randint(1, 5)
        average_score = sum([product_score_1, product_score_2, product_score_3,
                             service_score_1, service_score_2, service_score_3,
                             store_score_1, store_score_2, store_score_3]) / 9
        overall_score = randint_plusminus(math.floor(average_score), val_range=1)

        review = [company_name, review_id, employee_id, product_id,
                  product_score_1, product_score_2, product_score_3,
                  service_score_1, service_score_2, service_score_3,
                  store_score_1, store_score_2, store_score_3, overall_score]
        all_data.append(review)
        review_id += 1

        document.add_paragraph(f"{company_name}").style = Title_style
        document.add_paragraph(f"Product Review #{review[1]}\n")
        document.add_paragraph("Employee ID:\t").add_run(f"\t{review[2]}\t").underline = True
        document.add_paragraph("Product ID:\t").add_run(f"\t{review[3]}\t").underline = True

        table = document.add_table(rows=12, cols=6)
        table.style = 'Table Grid'
        for cell in table.columns[0].cells:
            cell.width = Inches(3)

        criterion1 = table.cell(0, 0)
        criterion2 = table.cell(1, 0)
        criterion_cell = criterion1.merge(criterion2)
        criterion_cell.text = 'Criterion'
        criterion_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        criterion_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        score1 = table.cell(0, 1)
        score2 = table.cell(0, 5)
        score_cell = score1.merge(score2)
        score_cell.text = 'Score'
        score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1, 1).text = '1'
        table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 2).text = '2'
        table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 3).text = '3'
        table.cell(1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 4).text = '4'
        table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 5).text = '5'
        table.cell(1, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2, 0).text = "The product is high quality"
        table.cell(3, 0).text = "The product has good packaging"
        table.cell(4, 0).text = "The product is easy to use"

        table.cell(5, 0).text = "The employee is polite"
        table.cell(6, 0).text = "The employee is helpful"
        table.cell(7, 0).text = "The employee is friendly"

        table.cell(8, 0).text = "The store is easy to navigate"
        table.cell(9, 0).text = "The store is clean"
        table.cell(10, 0).text = "The store has enough parking"

        table.cell(11, 0).text = "Overall satisfaction"

        for x in range(2, 12):
            table.cell(x, review[x + 2]).text = "X"
            table.cell(x, review[x + 2]).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    with open(f"review_csv/{company_name}.txt", "w+") as txt_file:
        for line in all_data:
            line = [str(_) for _ in line]
            txt_file.write(",".join(line) + "\n")

    document.save(f"review_docx/{company_name}.docx")
    print(review)
